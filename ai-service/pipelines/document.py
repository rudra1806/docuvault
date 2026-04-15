# ============================================================
# pipelines/document.py — Document Extraction Pipeline
# ============================================================
# Extracts text from PDF, DOCX, TXT, and DOC files.
# ============================================================

import io
import logging
import pdfplumber
from docx import Document as DocxDocument
import chardet

logger = logging.getLogger(__name__)


async def extract_document(file_bytes: bytes, file_name: str, ext: str) -> dict:
    """Extract text from document files (PDF, DOCX, TXT, DOC)."""

    if ext == "pdf":
        return _extract_pdf(file_bytes, file_name)
    elif ext == "docx":
        return _extract_docx(file_bytes, file_name)
    elif ext == "txt":
        return _extract_txt(file_bytes, file_name)
    elif ext == "doc":
        return _extract_doc(file_bytes, file_name)
    else:
        return {"text": "", "source": "document", "metadata": {"error": f"Unknown doc type: {ext}"}}


def _extract_pdf(file_bytes: bytes, file_name: str) -> dict:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    page_count = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""

            # Also try to extract table data
            tables = page.extract_tables()
            table_text = ""
            for table in tables:
                for row in table:
                    # Filter None values and join
                    row_clean = [str(cell) if cell else "" for cell in row]
                    table_text += " | ".join(row_clean) + "\n"

            combined = page_text
            if table_text:
                combined += f"\n[Table Data]\n{table_text}"

            if combined.strip():
                text_parts.append(f"[Page {i + 1}]\n{combined.strip()}")

    return {
        "text": "\n\n".join(text_parts),
        "source": "document",
        "metadata": {"pages": page_count, "type": "pdf"},
    }


def _extract_docx(file_bytes: bytes, file_name: str) -> dict:
    """Extract text from DOCX using python-docx."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract table data
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                text_parts.append(" | ".join(row_data))

    return {
        "text": "\n".join(text_parts),
        "source": "document",
        "metadata": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "type": "docx",
        },
    }


def _extract_txt(file_bytes: bytes, file_name: str) -> dict:
    """Extract text from plain text files with encoding detection."""
    # Detect encoding
    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        text = file_bytes.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        text = file_bytes.decode("utf-8", errors="replace")

    return {
        "text": text.strip(),
        "source": "document",
        "metadata": {"encoding": encoding, "type": "txt"},
    }


def _extract_doc(file_bytes: bytes, file_name: str) -> dict:
    """
    Extract text from legacy DOC files.
    Falls back to binary text extraction since antiword may not be available.
    """
    # Try basic binary text extraction
    try:
        detected = chardet.detect(file_bytes)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        text = file_bytes.decode(encoding, errors="replace")

        # Filter out non-printable characters but keep newlines and spaces
        cleaned = "".join(
            char for char in text
            if char.isprintable() or char in "\n\r\t"
        )

        return {
            "text": cleaned.strip(),
            "source": "document",
            "metadata": {"type": "doc", "note": "basic extraction"},
        }
    except Exception as e:
        logger.warning(f"DOC extraction failed for {file_name}: {e}")
        return {
            "text": "",
            "source": "document",
            "metadata": {"type": "doc", "error": str(e)},
        }
